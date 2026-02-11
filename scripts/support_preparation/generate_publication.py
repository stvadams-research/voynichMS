#!/usr/bin/env python3
"""
Voynich Manuscript Publication Generator (Data-Driven Pipeline)

Generates publication documents from authored markdown content, resolved against
real JSON data files and computed report markdown. Replaces all prior hardcoded
publication generators.

Architecture:
  Content .md files  ──┐
                       ├──→ [DataResolver] ──→ [Assembler] ──→ [DocxRenderer] ──→ .docx
  Data .json files   ──┘
  Report .md files   ──┘
  Visual .png files  ──┘

Usage:
  python generate_publication.py --profile full
  python generate_publication.py --profile summary
  python generate_publication.py --profile technical
  python generate_publication.py --list-missing   # show unresolved placeholders
"""

import re
import json
import argparse
from pathlib import Path
from datetime import datetime
import logging

import yaml
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
CONTENT_DIR = PROJECT_ROOT / "planning/support_preparation/content"
CONFIG_PATH = Path(__file__).resolve().parent / "publication_config.yaml"
OUTPUT_DIR = PROJECT_ROOT / "results/publication"


class DataResolver:
    """Resolves {{placeholder}} markers against JSON data files and static values."""

    def __init__(self, config: dict):
        self.sources = config.get("data_sources", {})
        self._cache: dict = {}
        self.missing: list[str] = []
        self.resolved: list[tuple[str, str, str]] = []  # (key, value, source)

    def _load_json(self, rel_path: str) -> dict:
        if rel_path not in self._cache:
            full_path = PROJECT_ROOT / rel_path
            if not full_path.exists():
                logger.warning(f"Data file not found: {full_path}")
                self._cache[rel_path] = None
            else:
                with open(full_path) as f:
                    self._cache[rel_path] = json.load(f)
        return self._cache[rel_path]

    def _navigate(self, data: dict, key_path: str):
        """Navigate nested dict/list using dot notation. Supports string keys with spaces."""
        parts = key_path.split(".")
        current = data
        for part in parts:
            if current is None:
                return None
            if isinstance(current, dict):
                if part in current:
                    current = current[part]
                else:
                    return None
            elif isinstance(current, list):
                try:
                    current = current[int(part)]
                except (ValueError, IndexError):
                    return None
            else:
                return None
        return current

    def resolve_key(self, key: str) -> str:
        """Resolve a single placeholder key to its string value."""
        if key not in self.sources:
            self.missing.append(key)
            return f"[MISSING: {key}]"

        source = self.sources[key]

        # Static string value
        if isinstance(source, str):
            self.resolved.append((key, source, "static"))
            return source

        # JSON file lookup
        if isinstance(source, dict):
            file_path = source.get("file", "")
            json_path = source.get("path", "")
            fmt = source.get("format", "")

            data = self._load_json(file_path)
            if data is None:
                self.missing.append(key)
                return f"[MISSING: {key} (file not found: {file_path})]"

            # Handle files that wrap results in a "results" key
            value = self._navigate(data, json_path)
            if value is None and "results" in data:
                value = self._navigate(data["results"], json_path)
            if value is None:
                # Try without "results" prefix removal
                self.missing.append(key)
                return f"[MISSING: {key} (path not found: {json_path} in {file_path})]"

            # Format
            if fmt == "percent" and isinstance(value, (int, float)):
                formatted = f"{value * 100:.2f}"
            elif isinstance(value, float):
                formatted = f"{value:.4f}"
            else:
                formatted = str(value)

            self.resolved.append((key, formatted, f"{file_path} -> {json_path}"))
            return formatted

        self.missing.append(key)
        return f"[MISSING: {key} (bad config)]"

    def resolve_text(self, text: str) -> str:
        """Resolve all {{placeholder}} markers in a text string."""
        def replacer(match):
            key = match.group(1).strip()
            return self.resolve_key(key)
        return re.sub(r"\{\{([^{}|]+?)\}\}", replacer, text)


class ReportIncluder:
    """Resolves {{include:path}} markers by reading the referenced markdown file."""

    def __init__(self):
        self.included: list[str] = []
        self.missing: list[str] = []

    def resolve(self, text: str) -> str:
        def replacer(match):
            rel_path = match.group(1).strip()
            full_path = PROJECT_ROOT / rel_path

            # Handle special :table suffix for JSON files
            if rel_path.endswith(":table"):
                return self._include_json_table(rel_path[:-6])

            if not full_path.exists():
                self.missing.append(rel_path)
                return f"\n\n[MISSING REPORT: {rel_path}]\n\n"

            with open(full_path) as f:
                content = f.read()
            self.included.append(rel_path)
            return f"\n\n{content}\n\n"

        return re.sub(r"\{\{include:(.+?)\}\}", replacer, text)

    def _include_json_table(self, rel_path: str) -> str:
        """Render a JSON file as a markdown table."""
        full_path = PROJECT_ROOT / rel_path
        if not full_path.exists():
            self.missing.append(rel_path)
            return f"\n\n[MISSING DATA: {rel_path}]\n\n"

        with open(full_path) as f:
            data = json.load(f)

        # Handle sectional_profiles.json style: dict of dicts
        if isinstance(data, dict) and all(isinstance(v, dict) for v in data.values()):
            rows = []
            # Collect all keys from sub-dicts
            sub_keys = set()
            for v in data.values():
                sub_keys.update(v.keys())
            sub_keys = sorted(sub_keys)

            header = "| Section | " + " | ".join(sub_keys) + " |"
            separator = "|" + "|".join(["---"] * (len(sub_keys) + 1)) + "|"
            rows.append(header)
            rows.append(separator)
            for section, values in sorted(data.items()):
                cells = [str(values.get(k, "")) for k in sub_keys]
                rows.append(f"| {section} | " + " | ".join(cells) + " |")

            self.included.append(rel_path)
            return "\n\n" + "\n".join(rows) + "\n\n"

        self.missing.append(rel_path)
        return f"\n\n[UNSUPPORTED JSON FORMAT: {rel_path}]\n\n"


class FigureCollector:
    """Collects {{figure:path|caption}} markers for later rendering."""

    def __init__(self, include_figures: bool = True):
        self.figures: list[tuple[str, str, str]] = []  # (marker_id, path, caption)
        self.missing: list[str] = []
        self.include_figures = include_figures
        self._counter = 0

    def resolve(self, text: str) -> str:
        def replacer(match):
            content = match.group(1).strip()
            parts = content.split("|", 1)
            rel_path = parts[0].strip()
            caption = parts[1].strip() if len(parts) > 1 else ""

            full_path = PROJECT_ROOT / rel_path
            if not full_path.exists():
                self.missing.append(rel_path)
                return f"\n\n[MISSING FIGURE: {rel_path}]\n\n"

            self._counter += 1
            marker = f"__FIGURE_{self._counter}__"
            self.figures.append((marker, str(full_path), caption))

            if not self.include_figures:
                return f"\n\n[Figure {self._counter}: {caption}]\n\n"
            return f"\n\n{marker}\n\n"

        return re.sub(r"\{\{figure:(.+?)\}\}", replacer, text)


class DocxRenderer:
    """Converts resolved markdown content to a Word document."""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)

        for level, (name, size) in enumerate(
            [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)], start=1
        ):
            h = self.doc.styles[name]
            h.font.name = "Arial"
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.space_before = Pt(20 if level < 3 else 12)
            h.paragraph_format.space_after = Pt(10 if level < 3 else 6)

        if "Title" in self.doc.styles:
            t = self.doc.styles["Title"]
        else:
            t = self.doc.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
        t.font.name = "Arial"
        t.font.size = Pt(22)
        t.font.bold = True
        t.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t.paragraph_format.space_after = Pt(24)

    def _add_formatted_text(self, paragraph, text: str):
        """Parse **bold** and *italic* markdown and add to paragraph."""
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                paragraph.add_run(part[2:-2]).bold = True
            elif part.startswith("*") and part.endswith("*"):
                paragraph.add_run(part[1:-1]).italic = True
            else:
                paragraph.add_run(part)

    def _is_table_row(self, line: str) -> bool:
        stripped = line.strip()
        return stripped.startswith("|") and stripped.endswith("|")

    def _is_separator_row(self, line: str) -> bool:
        if not self._is_table_row(line):
            return False
        cells = [cell.strip() for cell in line.split("|")[1:-1]]
        return all(set(c) <= {"-", ":", " "} for c in cells)

    def _render_table(self, lines: list[str]):
        rows = []
        for line in lines:
            if self._is_separator_row(line):
                continue
            cells = [cell.strip() for cell in line.split("|")[1:-1]]
            rows.append(cells)

        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                if j < num_cols:
                    p = table.cell(i, j).paragraphs[0]
                    self._add_formatted_text(p, cell_data)
                    if i == 0 and p.runs:
                        p.runs[0].bold = True
        self.doc.add_paragraph()

    def render(self, markdown: str, figures: list[tuple[str, str, str]]):
        """Render resolved markdown to the Word document."""
        figure_map = {marker: (path, caption) for marker, path, caption in figures}
        lines = markdown.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if not stripped:
                i += 1
                continue

            # Figure marker
            if stripped in figure_map:
                path, caption = figure_map[stripped]
                if Path(path).exists():
                    self.doc.add_picture(path, width=Inches(5.5))
                    cap_p = self.doc.add_paragraph(f"Figure: {caption}")
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cap_p.runs:
                        run.font.size = Pt(9)
                        run.font.italic = True
                i += 1
                continue

            # Headings
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                text = stripped.lstrip("#").strip()
                # Title-level heading (# with no preceding content = title page)
                if level == 1:
                    self.doc.add_heading(text, level=1)
                elif level == 2:
                    self.doc.add_heading(text, level=2)
                elif level >= 3:
                    self.doc.add_heading(text, level=3)
                i += 1
                continue

            # Horizontal rule = page break
            if stripped in ("---", "***", "___"):
                self.doc.add_page_break()
                i += 1
                continue

            # Table
            if self._is_table_row(stripped):
                table_lines = []
                while i < len(lines) and self._is_table_row(lines[i].strip()):
                    table_lines.append(lines[i])
                    i += 1
                self._render_table(table_lines)
                continue

            # [MISSING ...] markers - render in red
            if stripped.startswith("[MISSING"):
                p = self.doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.color.rgb = RGBColor(200, 0, 0)
                run.font.bold = True
                i += 1
                continue

            # Regular paragraph
            p = self.doc.add_paragraph()
            self._add_formatted_text(p, stripped)
            i += 1

    def save(self, path: Path):
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))


class PublicationBuilder:
    """Orchestrates the full pipeline: load config, resolve content, render document."""

    def __init__(self, profile: str = "full"):
        with open(CONFIG_PATH) as f:
            self.config = yaml.safe_load(f)

        if profile not in self.config["profiles"]:
            raise ValueError(
                f"Unknown profile '{profile}'. Available: {list(self.config['profiles'].keys())}"
            )

        self.profile = self.config["profiles"][profile]
        self.profile_name = profile
        self.resolver = DataResolver(self.config)
        self.includer = ReportIncluder()
        self.figures = FigureCollector(
            include_figures=self.profile.get("include_figures", True)
        )

    def _load_content(self) -> str:
        """Load and concatenate all chapter content files for this profile."""
        chapters = self.profile["chapters"]
        parts = []
        for chapter_name in chapters:
            path = CONTENT_DIR / f"{chapter_name}.md"
            if not path.exists():
                logger.warning(f"Content file not found: {path}")
                parts.append(f"\n\n[MISSING CONTENT FILE: {chapter_name}.md]\n\n")
                continue
            with open(path) as f:
                parts.append(f.read())
            parts.append("\n\n---\n\n")  # page break between chapters
        return "\n".join(parts)

    def build(self, list_missing: bool = False) -> Path:
        """Execute the full pipeline and return the output path."""
        logger.info(f"Building profile: {self.profile_name}")

        # Step 1: Load content
        raw_content = self._load_content()
        logger.info(
            f"Loaded {len(self.profile['chapters'])} content files"
        )

        # Step 2: Resolve report includes
        content = self.includer.resolve(raw_content)
        logger.info(
            f"Included {len(self.includer.included)} reports, "
            f"{len(self.includer.missing)} missing"
        )

        # Step 3: Resolve data placeholders
        content = self.resolver.resolve_text(content)
        logger.info(
            f"Resolved {len(self.resolver.resolved)} data values, "
            f"{len(self.resolver.missing)} missing"
        )

        # Step 4: Collect figure references
        content = self.figures.resolve(content)
        logger.info(
            f"Found {len(self.figures.figures)} figures, "
            f"{len(self.figures.missing)} missing"
        )

        if list_missing:
            self._print_missing_report()
            return None

        # Step 5: Render to docx
        renderer = DocxRenderer()
        renderer.render(content, self.figures.figures)

        output_path = OUTPUT_DIR / self.profile["output"]
        renderer.save(output_path)
        logger.info(f"Document saved: {output_path}")

        # Step 6: Write build log
        self._write_build_log(output_path)

        return output_path

    def _print_missing_report(self):
        """Print all unresolved references."""
        all_missing = []
        if self.resolver.missing:
            print("\n=== Missing Data Placeholders ===")
            for key in sorted(set(self.resolver.missing)):
                print(f"  {{{{ {key} }}}}")
                all_missing.append(f"data: {key}")

        if self.includer.missing:
            print("\n=== Missing Report Includes ===")
            for path in sorted(set(self.includer.missing)):
                print(f"  {{{{include:{path}}}}}")
                all_missing.append(f"report: {path}")

        if self.figures.missing:
            print("\n=== Missing Figures ===")
            for path in sorted(set(self.figures.missing)):
                print(f"  {{{{figure:{path}}}}}")
                all_missing.append(f"figure: {path}")

        if not all_missing:
            print("\nAll references resolved successfully.")
        else:
            print(f"\nTotal missing: {len(all_missing)}")

    def _write_build_log(self, output_path: Path):
        """Write a provenance log alongside the output document."""
        log_path = output_path.with_suffix(".build_log.txt")
        with open(log_path, "w") as f:
            f.write(f"Publication Build Log\n")
            f.write(f"Generated: {datetime.now().isoformat()}\n")
            f.write(f"Profile: {self.profile_name}\n")
            f.write(f"Output: {output_path}\n")
            f.write(f"\n{'='*60}\n")

            f.write(f"\nData Sources Resolved ({len(self.resolver.resolved)}):\n")
            for key, value, source in sorted(self.resolver.resolved):
                f.write(f"  {key} = {value}  [{source}]\n")

            if self.resolver.missing:
                f.write(f"\nMissing Data ({len(self.resolver.missing)}):\n")
                for key in sorted(set(self.resolver.missing)):
                    f.write(f"  {key}\n")

            f.write(f"\nReports Included ({len(self.includer.included)}):\n")
            for path in sorted(self.includer.included):
                f.write(f"  {path}\n")

            if self.includer.missing:
                f.write(f"\nMissing Reports ({len(self.includer.missing)}):\n")
                for path in sorted(set(self.includer.missing)):
                    f.write(f"  {path}\n")

            f.write(f"\nFigures Embedded ({len(self.figures.figures)}):\n")
            for marker, path, caption in self.figures.figures:
                f.write(f"  {Path(path).name}: {caption}\n")

            if self.figures.missing:
                f.write(f"\nMissing Figures ({len(self.figures.missing)}):\n")
                for path in sorted(set(self.figures.missing)):
                    f.write(f"  {path}\n")

        logger.info(f"Build log saved: {log_path}")


def main():
    parser = argparse.ArgumentParser(
        description="Voynich Publication Generator (Data-Driven Pipeline)"
    )
    parser.add_argument(
        "--profile",
        default="full",
        help="Output profile: full, summary, technical (default: full)",
    )
    parser.add_argument(
        "--list-missing",
        action="store_true",
        help="List all unresolved placeholders without generating output",
    )
    args = parser.parse_args()

    builder = PublicationBuilder(profile=args.profile)
    result = builder.build(list_missing=args.list_missing)

    if result:
        print(f"\nGenerated: {result}")


if __name__ == "__main__":
    main()
